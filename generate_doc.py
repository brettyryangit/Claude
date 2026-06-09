from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ── Colour palette ──────────────────────────────────────────────────────────
BLACK     = RGBColor(0x1a, 0x1a, 0x1a)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MID_GRAY  = RGBColor(0x55, 0x55, 0x55)
LIGHT_BG  = RGBColor(0xF2, 0xF2, 0xF2)
ACCENT    = RGBColor(0xC0, 0x39, 0x2B)   # red accent for key callouts
TABLE_HDR = RGBColor(0x1a, 0x1a, 0x1a)

# ── Helper: shade a table cell ───────────────────────────────────────────────
def shade_cell(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

# ── Helper: set paragraph border (top rule) ──────────────────────────────────
def add_top_border(para, color_hex='C0392B', size=12):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    str(size))
    top.set(qn('w:space'), '4')
    top.set(qn('w:color'), color_hex)
    pBdr.append(top)
    pPr.append(pBdr)

# ── Helper: bold run ─────────────────────────────────────────────────────────
def bold_run(para, text, size=11, color=None):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

# ═══════════════════════════════════════════════════════════════════════════
#  COVER / TITLE BLOCK
# ═══════════════════════════════════════════════════════════════════════════
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_top_border(title_para, 'C0392B', 24)
r = title_para.add_run('TRADING PLAN')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = BLACK

sub = doc.add_paragraph()
r2 = sub.add_run('Scalping ES / NQ & Crypto  ·  1m – 5m Timeframes')
r2.font.size = Pt(13)
r2.font.color.rgb = MID_GRAY
r2.italic = True

note = doc.add_paragraph()
r3 = note.add_run(
    'This document exists for one reason: to keep you on script. '
    'If a trade does not check every box in this plan, it does not exist.'
)
r3.font.size = Pt(11)
r3.font.color.rgb = ACCENT
r3.bold = True

doc.add_paragraph()   # spacer

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION HELPER
# ═══════════════════════════════════════════════════════════════════════════
def section_heading(title):
    p = doc.add_paragraph()
    add_top_border(p, '1a1a1a', 6)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = BLACK
    return p

def sub_heading(title):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = DARK_GRAY
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = DARK_GRAY
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = DARK_GRAY
    return p

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = DARK_GRAY
    return p

def spacer():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return p

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 1 — THE PROBLEM
# ═══════════════════════════════════════════════════════════════════════════
section_heading('The Problem This Plan Solves')
body(
    'You already know how to trade. The issue is impulsive, off-plan entries — '
    'chasing moves, clicking without confirmation, trading on your phone, acting '
    'on feelings rather than structure. This plan is a hard ruleset, not a '
    'guideline. Rules only work if you treat exceptions as failures, not '
    'opportunities.'
)
spacer()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 2 — MARKETS & SESSIONS
# ═══════════════════════════════════════════════════════════════════════════
section_heading('Section 1 — Markets & Sessions')

sub_heading('Instruments')
bullet('Futures: ES (S&P 500), NQ (Nasdaq 100)')
bullet('Crypto: BTC/USD, ETH/USD  (expand only after 30 consecutive on-plan days)')
spacer()

sub_heading('Valid Trading Sessions')
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
hdr = tbl.rows[0].cells
for i, txt in enumerate(['Session', 'ES / NQ', 'Crypto']):
    shade_cell(hdr[i], '1a1a1a')
    p = hdr[i].paragraphs[0]
    r = p.add_run(txt)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(10)

rows_data = [
    ('Pre-market  8:30 – 9:30 AM ET',      '✓', 'High-volume hours only'),
    ('New York Open  9:30 – 11:30 AM ET',   '✓', '✓'),
    ('London Close / NY  10:00 AM – 12:00 PM ET', '✓', '✓'),
    ('Afternoon  1:30 – 3:30 PM ET',        '✓', '✓'),
    ('AVOID  11:30 AM – 1:30 PM ET (chop)', '✗', 'Low-volume dead hours'),
]
for i, (s, f, c) in enumerate(rows_data):
    row = tbl.add_row().cells
    fill = 'F2F2F2' if i % 2 == 0 else 'FFFFFF'
    for cell in row:
        shade_cell(cell, fill)
    for j, txt in enumerate([s, f, c]):
        p = row[j].paragraphs[0]
        r = p.add_run(txt)
        r.font.size = Pt(10)
        if txt == '✗':
            r.font.color.rgb = ACCENT
        else:
            r.font.color.rgb = DARK_GRAY

spacer()
p = doc.add_paragraph()
r = p.add_run('Rule: ')
r.bold = True; r.font.color.rgb = ACCENT; r.font.size = Pt(10.5)
r2 = p.add_run('If the clock is not in a valid session window, there is no trade. Close the chart.')
r2.font.size = Pt(10.5); r2.font.color.rgb = DARK_GRAY
spacer()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 3 — REQUIRED SETUP CRITERIA
# ═══════════════════════════════════════════════════════════════════════════
section_heading('Section 2 — Required Setup Criteria (All 5 Required)')
body(
    'Every single entry requires all five criteria to be confirmed before '
    'touching the buy/sell button. Not four. Not three. All five.'
)
spacer()

criteria = [
    (
        '1.  HTF Bias Is Clear',
        [
            'Check the 15m or 1h chart first.',
            'Identify whether price is in a premium (shorting zone) or discount (buying zone) relative to the most recent swing range.',
            'Identify the nearest Fair Value Gaps (FVG) and Order Blocks (OB) on HTF.',
            'You must be able to state your bias in one sentence. If you can\'t, there is no trade.',
        ]
    ),
    (
        '2.  Change of Structure (CHoCH / BOS) on Entry Timeframe',
        [
            'On the 1m or 5m chart, a Break of Structure (BOS) or Change of Character (CHoCH) must have already printed — not be forming, not be guessed at.',
            'BOS = continuation of trend structure (momentum entry).',
            'CHoCH = reversal signal (counter-trend, higher caution required).',
            'No structure break = no entry, no exceptions.',
        ]
    ),
    (
        '3.  Liquidity Grab Preceding the Move',
        [
            'A stop hunt or liquidity sweep must have occurred before your entry.',
            'Price swept above a recent high (for a short) or below a recent low (for a long), then reversed.',
            'Look for equal highs/lows as liquidity targets.',
            'If there was no liquidity grab, you are not entering the move — you are chasing it.',
        ]
    ),
    (
        '4.  Volume Confirmation',
        [
            'The move into your entry zone must show increased relative volume.',
            'For reversals: volume spike on the sweep candle + decreasing volume on the pullback.',
            'For continuations: volume expands on the BOS candle.',
            'A low-volume move into a level is a trap, not a setup.',
        ]
    ),
    (
        '5.  Point of Interest (POI) Alignment',
        [
            'Entry must be at or within a defined POI: Order Block, Fair Value Gap, Breaker Block, or Mitigation Block.',
            'The POI must be identified before price reaches it, not retrofitted after.',
            'Chasing a candle that already ran into a level is not an entry.',
        ]
    ),
]

for title, points in criteria:
    sub_heading(title)
    for pt in points:
        bullet(pt)
    spacer()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 4 — ENTRY PROTOCOL
# ═══════════════════════════════════════════════════════════════════════════
section_heading('Section 3 — Entry Protocol (Step by Step)')
body('Do not skip steps. Do not reorder steps. Run this process every time.')
spacer()

steps = [
    ('Step 1 — Check the clock.',
     'Are you in a valid session? If no → close the chart, walk away.'),
    ('Step 2 — Read HTF bias.',
     '15m or 1h: where is price? Premium or discount? What are the nearest HTF POIs? Write it down or say it out loud.'),
    ('Step 3 — Identify the setup on 1m / 5m.',
     'Has a liquidity grab occurred? Has a CHoCH or BOS printed? Is price now retracing to a POI?'),
    ('Step 4 — Confirm volume.',
     'Is volume confirming the structural move?'),
    ('Step 5 — Define the trade BEFORE entering.',
     'Entry price / Stop loss (beyond the liquidity sweep candle) / Target (next HTF POI) / R:R (minimum 1:2) — all four must be written before you touch the button.'),
    ('Step 6 — Override test.',
     '"Would I take this trade in a funded account evaluation?" If no → do not take it.'),
    ('Step 7 — Enter.',
     'Only if Steps 1–6 are all clear.'),
]

for label, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r1 = p.add_run(label + '  ')
    r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = BLACK
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5); r2.font.color.rgb = DARK_GRAY

spacer()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 5 — RISK MANAGEMENT
# ═══════════════════════════════════════════════════════════════════════════
section_heading('Section 4 — Risk Management')

sub_heading('Position Sizing')
bullet('Maximum risk per trade: 1% of account balance.')
bullet('Maximum daily loss: 2% of account balance. Hit it → session over, charts closed.')
bullet('Maximum concurrent trades: 1 (scalping — one thing at a time).')
spacer()

sub_heading('Stop Loss Rules')
bullet('Stop is always placed beyond the liquidity sweep candle — not at a round number, not "tight to manage risk."')
bullet('Never move stop loss to a worse position (never widen a stop).')
bullet('Moving stop to break-even is allowed once price moves 1:1 toward target.')
spacer()

sub_heading('Take Profit Rules')
bullet('Minimum target is 1:2 risk/reward. Below this, the trade does not meet criteria.')
bullet('Partial profit at 1:1 (move stop to BE), remainder to full target is acceptable.')
bullet('Do not move target based on "feeling." The target is set in Step 5 and does not change.')
spacer()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 6 — BEHAVIOURAL RULES
# ═══════════════════════════════════════════════════════════════════════════
section_heading('Section 5 — Behavioral Rules (Non-Negotiable)')
body('These rules exist specifically because of the impulsive trading pattern.')
spacer()

sub_heading('Device Rules')
bullet('No trading on the phone. Period. The phone is for monitoring positions already open — never for entering new ones.')
bullet('Charts are analyzed on the computer only. If you are not at your computer, you are not trading.')
spacer()

sub_heading('Pre-Market Prep (Required Before Opening the Platform)')
bullet('Mark HTF POIs on the chart before the session — not during a live trade.')
bullet('Identify today\'s likely liquidity targets: equal highs/lows, overnight high/low.')
bullet('Write your bias down.')
bullet('If you have not done prep, you are not ready to trade. Do not open the platform.')
spacer()

sub_heading('Emotional State Check')
bullet('Before each session, rate your mental state 1–10.')
bullet('Score 6 or below (stressed, frustrated, distracted, tired) → do not trade that session.')
bullet('Revenge trading is always a 6 or below situation. Do not fool yourself.')
spacer()

sub_heading('After a Loss')
bullet('Take a 15-minute break away from charts after any losing trade.')
bullet('Do not immediately re-enter to "make it back." This is the single most expensive scalping habit.')
bullet('After 2 consecutive losses, the session is over — regardless of where you are in the daily loss limit.')
spacer()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 7 — ON-PLAN vs OFF-PLAN
# ═══════════════════════════════════════════════════════════════════════════
section_heading('Section 6 — On-Plan vs Off-Plan Trade Examples')

sub_heading('What an On-Plan Trade Looks Like')
steps_onplan = [
    'HTF shows price at a discount (for longs) or premium (for shorts).',
    'Price sweeps a liquidity level (equal lows for longs, equal highs for shorts).',
    'A CHoCH or BOS prints on the 1m/5m after the sweep.',
    'Price pulls back into the nearest FVG or OB created by the displacement move.',
    'Volume on the displacement candle is elevated.',
    'You enter at the POI with stop beyond the sweep candle.',
    'Target is the next HTF POI — minimum 1:2.',
]
for s in steps_onplan:
    numbered(s)
spacer()

sub_heading('What an Off-Plan Trade Looks Like')
body('Know these patterns — you have taken all of them.')
off_plan = [
    'Entering because "it looks like it\'s going up."',
    'Chasing a candle that already moved without waiting for a pullback.',
    'Trading during dead hours (11:30–1:30 ET) because "it\'s moving."',
    'Trading on the phone because "I see it and don\'t want to miss it."',
    'Adding to a losing position to average down.',
    'Entering without a defined stop.',
    'Entering with less than 1:2 R:R because "it\'s a quick scalp."',
    'Trading a second position before the first one is closed.',
    'Taking a trade immediately after a loss without the 15-minute break.',
]
for item in off_plan:
    bullet(item)

p = doc.add_paragraph()
r = p.add_run('If the trade you are about to take matches anything above, close the order ticket.')
r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = ACCENT
spacer()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 8 — DAILY REVIEW
# ═══════════════════════════════════════════════════════════════════════════
section_heading('Section 7 — Daily Review Process')
body('10–15 minutes after every session. Not optional. Fill this out for every trade taken.')
spacer()

review_rows = [
    ('Was it a valid session time?',              'Y / N'),
    ('Was HTF bias identified before entry?',     'Y / N'),
    ('Was there a liquidity grab?',               'Y / N'),
    ('Was there a CHoCH or BOS?',                 'Y / N'),
    ('Was volume confirming?',                    'Y / N'),
    ('Was the entry at a pre-defined POI?',       'Y / N'),
    ('Was R:R at least 1:2?',                     'Y / N'),
    ('Was stop loss properly placed?',            'Y / N'),
    ('ON-PLAN?',                                  'Y / N'),
]
tbl2 = doc.add_table(rows=1, cols=2)
tbl2.style = 'Table Grid'
hdr2 = tbl2.rows[0].cells
for i, txt in enumerate(['Question', 'Answer']):
    shade_cell(hdr2[i], '1a1a1a')
    p = hdr2[i].paragraphs[0]
    r = p.add_run(txt)
    r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); r.font.size = Pt(10)

for i, (q, a) in enumerate(review_rows):
    row = tbl2.add_row().cells
    fill = 'F2F2F2' if i % 2 == 0 else 'FFFFFF'
    is_last = (i == len(review_rows) - 1)
    if is_last:
        fill = '1a1a1a'
    for cell in row:
        shade_cell(cell, fill)
    for j, txt in enumerate([q, a]):
        p = row[j].paragraphs[0]
        r = p.add_run(txt)
        r.font.size = Pt(10)
        r.bold = is_last
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF) if is_last else DARK_GRAY

spacer()
body('Track your on-plan % daily. The goal is consistency, not P&L. Consistently on-plan → P&L follows.')
spacer()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 9 — PRE-TRADE CHECKLIST
# ═══════════════════════════════════════════════════════════════════════════
section_heading('Section 8 — Pre-Trade Checklist  (Print & Keep Visible)')
body('Run this before every entry. All 10 checked = trade is valid. Anything unchecked = no trade.')
spacer()

checklist = [
    'Clock — am I in a valid session?',
    'HTF bias — stated out loud or written down?',
    'Liquidity grab — did price sweep a high or low first?',
    'Structure break — CHoCH or BOS printed on 1m/5m?',
    'Volume — is it confirming the move?',
    'POI — am I entering at a defined level, not chasing?',
    'Entry defined — do I have a specific entry price?',
    'Stop defined — is it placed beyond the sweep candle?',
    'Target defined — is R:R at least 1:2?',
    'State check — am I a 7/10 or better mentally? On my computer, not my phone?',
]
for item in checklist:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r1 = p.add_run('☐  ')
    r1.font.size = Pt(12); r1.font.color.rgb = ACCENT; r1.bold = True
    r2 = p.add_run(item)
    r2.font.size = Pt(10.5); r2.font.color.rgb = DARK_GRAY

spacer()

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 10 — PSYCHOLOGY
# ═══════════════════════════════════════════════════════════════════════════
section_heading('Section 9 — Why You Go Off Script (and What to Do About It)')
body(
    'The pattern: you see movement, dopamine fires, impulse overrides process. '
    'This is not a trading problem — it\'s a behavioral loop. The trade feels urgent. '
    'Missing it feels like a loss. But entering without a setup is the loss.'
)
spacer()

p = doc.add_paragraph()
r = p.add_run('Reframe: ')
r.bold = True; r.font.color.rgb = ACCENT; r.font.size = Pt(10.5)
r2 = p.add_run(
    'The discipline trade is the one you don\'t take. Every time you see a move '
    'that doesn\'t meet criteria and you stay out, that is a won trade. Log it. '
    'Track your "no-trade" decisions. Reward the process.'
)
r2.font.size = Pt(10.5); r2.font.color.rgb = DARK_GRAY
spacer()

sub_heading('Practical Tools')
bullet('Keep the checklist physical and visible at your desk — not digital.')
bullet('Before clicking, say the setup out loud. Verbalizing forces deliberate thinking to engage.')
bullet('Rule: if you can\'t explain the trade in 30 seconds using your criteria, it\'s not a trade.')
bullet('If you feel the urge to trade outside your rules, close the platform for 10 minutes.')
spacer()

# ═══════════════════════════════════════════════════════════════════════════
#  FOOTER
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
add_top_border(p, 'C0392B', 6)
r = p.add_run('Last updated: 2026-06-09')
r.font.size = Pt(9); r.font.color.rgb = MID_GRAY; r.italic = True

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save('/home/user/Claude/Trading_Plan.docx')
print('Done — Trading_Plan.docx written.')
